from flask import Flask, render_template, request, make_response, session
import random
import io
from docx import Document

app = Flask(__name__)
app.secret_key = 'your_secret_key'

DAYS = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday"]
def insert_lunch_break_column(time_slots, lunch_break):
    """
    Inserts the lunch break into the time slots at the correct position dynamically.
    """
    updated_slots = []
    for i, slot in enumerate(time_slots):
        updated_slots.append(slot)
        if lunch_break in slot:  # Add lunch break column after the matching slot
            updated_slots.append(lunch_break)
    if lunch_break not in updated_slots:  # Handle edge case for lunch break not matching
        mid_index = len(time_slots) // 2
        updated_slots = time_slots[:mid_index + 1] + [lunch_break] + time_slots[mid_index + 1:]
    return updated_slots
# Helper function to insert lunch break as a column after the specified time slot
def insert_lunch_break_column(time_slots, lunch_break):
    index = next((i for i, slot in enumerate(time_slots) if lunch_break in slot), len(time_slots) // 2)
    updated_slots = time_slots[:index + 1] + [lunch_break] + time_slots[index + 1:]
    return updated_slots

# Function to generate the timetable
def generate_timetable(teachers, teacher_subject_map, subjects, time_slots, classrooms, lunch_break):
    # Insert lunch break into time slots dynamically
    time_slots_with_break = insert_lunch_break_column(time_slots, lunch_break)

    timetable = {day: {slot: None for slot in time_slots_with_break} for day in DAYS}

    for day in DAYS:
        for i, slot in enumerate(time_slots_with_break):
            if slot == lunch_break:  # Skip the lunch break slot
                timetable[day][slot] = "Lunch Break"
                continue

            available_subjects = [subj for subj in subjects if subj["hours"] > 0]
            if not available_subjects:
                continue

            for _ in range(10):  # Retry mechanism for constraints
                subject = random.choice(available_subjects)
                teacher_candidates = [
                    t for t in teachers
                    if subject["name"] in teacher_subject_map.get(t, [])
                ]

                if teacher_candidates:
                    teacher = random.choice(teacher_candidates)

                    # Check if the subject requires continuous classes
                    if subject["continuous_count"] == 1:
                        if (
                            i + 1 < len(time_slots_with_break)  # Ensure there's a next slot
                            and time_slots_with_break[i + 1] != lunch_break  # Avoid spanning lunch
                            and not timetable[day][time_slots_with_break[i]]  # Current slot free
                            and not timetable[day][time_slots_with_break[i + 1]]  # Next slot free
                        ):
                            # Assign two consecutive slots
                            timetable[day][time_slots_with_break[i]] = f"{subject['name']} - {teacher}"
                            timetable[day][time_slots_with_break[i + 1]] = f"{subject['name']} - {teacher}"
                            subject["hours"] -= 2  # Deduct two hours
                            break
                    else:
                        # Assign a single slot if no continuous class is required
                        if not timetable[day][time_slots_with_break[i]]:
                            timetable[day][time_slots_with_break[i]] = f"{subject['name']} - {teacher}"
                            subject["hours"] -= 1
                            break

    return timetable

# Route to generate timetable
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        teachers = request.form["teachers"].split(",")
        time_slots = request.form["time_slots"].split(",")
        lunch_break = request.form["lunch_break"].strip()
        classrooms = request.form["classrooms"].split(",")

        subjects = []
        for i in range(len(request.form.getlist("subject_name"))):
            subjects.append({
                "name": request.form.getlist("subject_name")[i].strip(),
                "hours": int(request.form.getlist("subject_hours")[i]),
                "is_lab": request.form.getlist("is_lab")[i].lower() == "yes",
                "continuous_count": int(request.form.getlist("continuous_count")[i])
            })

        teacher_subject_map = {}
        for i in range(len(request.form.getlist("teacher_name"))):
            teacher_subject_map[request.form.getlist("teacher_name")[i].strip()] = request.form.getlist("teacher_subjects")[i].split(",")

        timetable = generate_timetable(
            teachers, teacher_subject_map, subjects, time_slots, classrooms, lunch_break
        )

        session['timetable'] = timetable  # Save timetable to session
        session['time_slots'] = time_slots
        return render_template(
            "timetable.html", timetable=timetable, time_slots=time_slots, days=DAYS
        )

    return render_template("index.html")

# Route to download timetable as Word
@app.route("/download/word")
def download_word():
    timetable = session.get('timetable')  # Retrieve timetable from session
    time_slots = session.get('time_slots')  # Retrieve time slots from session
    if not timetable or not time_slots:
        return "Error: Timetable not available", 400

    document = Document()
    document.add_heading("Generated Timetable", level=1)

    # Create table with time slots as headings
    table = document.add_table(rows=1, cols=len(time_slots) + 1)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Day"
    for i, slot in enumerate(time_slots):
        header_cells[i + 1].text = slot

    # Fill in timetable data
    for day, slots in timetable.items():
        row_cells = table.add_row().cells
        row_cells[0].text = day
        for i, slot in enumerate(time_slots):
            row_cells[i + 1].text = slots[slot] or "Free"

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    response = make_response(file_stream.getvalue())
    response.headers["Content-Disposition"] = "attachment; filename=timetable.docx"
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return response

if __name__ == "__main__":
    app.run(debug=True)
